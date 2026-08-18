from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(r"E:\DATTTN\TTTN")
DOCX = ROOT / "DamMinhLinh_ A03_BCDK1_ChuongII_hoanthien.docx"
SOURCE = ROOT / "scripts" / "complete_chapter2.ps1"


def set_run_font(run, name="Times New Roman", size=13, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def insert_paragraph_before(anchor, text, style="Normal", code=False, caption=False):
    p = OxmlElement("w:p")
    anchor._p.addprevious(p)
    paragraph = anchor._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    p.addnext(paragraph._p)
    paragraph.text = text
    paragraph.style = style
    paragraph.paragraph_format.space_after = Pt(3)
    if caption:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=13, bold=True)
    elif code:
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            set_run_font(run, "Courier New", 10)
    elif style == "Heading 2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(6)
        for run in paragraph.runs:
            set_run_font(run, size=13, bold=True)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0.5)
        paragraph.paragraph_format.line_spacing = 1.3
        for run in paragraph.runs:
            set_run_font(run, size=13)
    return paragraph


def insert_table_before(anchor, rows, cols):
    table = anchor._parent.add_table(rows=rows, cols=cols, width=Cm(15.2))
    table.style = "Table Grid"
    anchor._p.addprevious(table._tbl)
    return table


def find_paragraph(doc, exact):
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise RuntimeError(f"Không tìm thấy đoạn: {exact}")


def remove_between(start, end):
    node = start._p.getnext()
    while node is not None and node is not end._p:
        nxt = node.getnext()
        if node.tag == qn("w:p"):
            node.getparent().remove(node)
        node = nxt


def parse_features(source):
    pattern = re.compile(r"@\('([^']+)','([^']+)','([^']+)'\)")
    values = pattern.findall(source)
    if len(values) != 54:
        raise RuntimeError(f"Kỳ vọng 54 đặc trưng, đọc được {len(values)}")
    return values


def style_table(table, features):
    headers = ["STT", "Tên đặc trưng", "Đơn vị", "Ý nghĩa/cách tính"]
    for j, value in enumerate(headers):
        table.cell(0, j).text = value
    for i, (name, unit, meaning) in enumerate(features, 1):
        vals = [str(i), name, unit, meaning]
        for j, value in enumerate(vals):
            table.cell(i, j).text = value
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    widths = [Cm(1.0), Cm(5.2), Cm(2.0), Cm(7.0)]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = widths[col_idx]
            cell.vertical_alignment = 1
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=9, bold=(row_idx == 0))


def add_reference_before(anchor, text):
    p = insert_paragraph_before(anchor, text)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)


def main():
    source = SOURCE.read_text(encoding="utf-8-sig")
    features = parse_features(source)
    doc = Document(DOCX)
    chapter2 = find_paragraph(doc, "CHƯƠNG 2. MÔ HÌNH ĐỀ XUẤT")
    chapter3 = find_paragraph(doc, "CHƯƠNG 3. KẾT QUẢ THỰC NGHIỆM VÀ THẢO LUẬN")
    appendix = find_paragraph(doc, "PHỤ LỤC")
    remove_between(chapter2, chapter3)

    table_added = False
    for line in source.splitlines():
        m = re.match(r"Add-Para '(.+)'(?: Heading2)?$", line)
        if m:
            text = m.group(1).replace("''", "'")
            style = "Heading 2" if line.endswith(" Heading2") else "Normal"
            insert_paragraph_before(chapter3, text, style=style)
            continue
        m = re.match(r'Add-Code "(.*)"$', line)
        if m:
            text = m.group(1).replace("`n", "\n").replace("`$", "$")
            insert_paragraph_before(chapter3, text, code=True)
            continue
        if not table_added and "Bảng 2.1. Danh sách 54 đặc trưng" in line:
            insert_paragraph_before(
                chapter3,
                "Bảng 2.1. Danh sách 54 đặc trưng của nids.flow_features.v1",
                caption=True,
            )
            table = insert_table_before(chapter3, 55, 4)
            style_table(table, features)
            table_added = True

    for text in re.findall(r"Add-Ref '(.+)'", source):
        add_reference_before(appendix, text.replace("''", "'"))

    doc.save(DOCX)


if __name__ == "__main__":
    main()
