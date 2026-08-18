from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

path = Path(r"E:\DATTTN\TTTN\DamMinhLinh_ A03_BCDK1_ChuongII_hoanthien_v2.docx")
doc = Document(path)


def find_paragraph(text):
    return next(p for p in doc.paragraphs if p.text.strip() == text)


chapter2 = find_paragraph("CHƯƠNG 2. MÔ HÌNH ĐỀ XUẤT")
chapter3 = find_paragraph("CHƯƠNG 3. KẾT QUẢ THỰC NGHIỆM VÀ THẢO LUẬN")

# Bỏ toàn bộ khối code/lệnh chạy trong Chương II.
node = chapter2._p.getnext()
while node is not None and node is not chapter3._p:
    nxt = node.getnext()
    if node.tag == qn("w:p"):
        text = "".join(t.text or "" for t in node.findall(".//" + qn("w:t"))).strip()
        is_code = any(
            (fonts.get(qn("w:ascii")) == "Courier New" or fonts.get(qn("w:hAnsi")) == "Courier New")
            for fonts in node.findall(".//" + qn("w:rFonts"))
        )
        command_prefixes = (
            "ip -br link", "sudo modprobe vfio-pci", "sudo python3 $DPDK_ROOT",
            "python3 scripts/dpdk_smoke.py", "sudo python3 scripts/dpdk_passive_probe.py",
            "python -m nids_mvp.preprocessing",
        )
        if is_code or text.startswith(command_prefixes):
            node.getparent().remove(node)
    node = nxt

# Chỉnh các câu phụ thuộc vào khối lệnh đã bỏ và lược phần hướng dẫn chạy chi tiết.
replacements = {
    "Hai lệnh ip xác nhận tên giao diện, trạng thái link và default route. ethtool -i ens160 xác nhận driver gốc vmxnet3. Hai liên kết sysfs cho biết địa chỉ PCI và IOMMU group. Script dpdk_smoke.py không chấp nhận cấu hình nếu thiếu IOMMU, nếu card dữ liệu trùng card quản trị, nếu group không phù hợp chính sách an toàn hoặc nếu đường quản trị không thể được bảo toàn. Dự án không bật chế độ VFIO no-IOMMU vì config khóa require_iommu=true và allow_no_iommu=false.":
        "Trước khi binding, hệ thống xác nhận tên giao diện, trạng thái link, default route, driver gốc vmxnet3, địa chỉ PCI và IOMMU group. Script dpdk_smoke.py không chấp nhận cấu hình nếu thiếu IOMMU, nếu card dữ liệu trùng card quản trị, nếu group không phù hợp chính sách an toàn hoặc nếu đường quản trị không thể được bảo toàn. Dự án không bật chế độ VFIO no-IOMMU vì config khóa require_iommu=true và allow_no_iommu=false.",
    "Các câu lệnh trên minh họa đúng thao tác mà workflow thực hiện có kiểm soát. Giá trị thực tế phải được đọc lại sau khi ghi; nếu kernel không cấp đủ 128 trang, script dừng thay vì tiếp tục với cấu hình khác. Documentation DPDK xác nhận /dev/hugepages là mount point thông dụng trên Linux hiện đại [1].":
        "Workflow đọc lại số hugepage sau khi cấu hình; nếu kernel không cấp đủ 128 trang, quy trình dừng thay vì tiếp tục với trạng thái không đúng hợp đồng. Documentation DPDK xác nhận /dev/hugepages là mount point thông dụng trên Linux hiện đại [1].",
    "discover chỉ thu thập và ghi bằng chứng; preflight kiểm tra topology, toolchain, devbind, hugepage và IOMMU; apply mới cấp hugepage và binding; run gọi dpdk-testpmd với EAL/application arguments đã khóa; rollback phục hồi host. Công cụ kiểm tra receipt từ chối bằng chứng thiếu bước hoặc sai cấu hình. testpmd cung cấp chế độ nhận và thống kê port; các lệnh show port stats/xstats được DPDK mô tả chính thức [6].":
        "Quy trình smoke test gồm kiểm tra điều kiện đầu vào, áp dụng cấu hình tạm thời, chạy testpmd trong thời gian giới hạn, thu thập thống kê và phục hồi host. Receipt chỉ được chấp nhận khi đầy đủ bằng chứng và đúng cấu hình. Các thống kê port của testpmd được DPDK mô tả chính thức [6].",
}
for p in doc.paragraphs:
    if p.text in replacements:
        p.text = replacements[p.text]
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = doc.styles["Normal"].font.size
            run.font.color.rgb = None
            color = run._element.get_or_add_rPr().find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                run._element.get_or_add_rPr().insert(1, color)
            color.set(qn("w:val"), "000000")


def set_header_refs(sectpr, rid):
    for ref in list(sectpr.findall(qn("w:headerReference"))):
        sectpr.remove(ref)
    default_ref = OxmlElement("w:headerReference")
    default_ref.set(qn("w:type"), "default")
    default_ref.set(qn("r:id"), rid)
    first_ref = OxmlElement("w:headerReference")
    first_ref.set(qn("w:type"), "first")
    first_ref.set(qn("r:id"), rid)
    sectpr.insert(0, first_ref)
    sectpr.insert(0, default_ref)


# Sau khi tách Chương II, section 18 là Tài liệu tham khảo và section 19 là Phụ lục.
section_properties = []
for element in doc._element.body:
    ppr = element.find(qn("w:pPr")) if element.tag == qn("w:p") else None
    sectpr = ppr.find(qn("w:sectPr")) if ppr is not None else None
    if sectpr is not None:
        section_properties.append(sectpr)
section_properties.append(doc._element.body.sectPr)

if len(section_properties) != 19:
    raise RuntimeError(f"Kỳ vọng 19 section, thực tế {len(section_properties)}")
set_header_refs(section_properties[17], "rId33")  # TÀI LIỆU THAM KHẢO
set_header_refs(section_properties[18], "rId34")  # PHỤ LỤC

doc.save(path)
