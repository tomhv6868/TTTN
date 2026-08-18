from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

path = Path(r"E:\DATTTN\TTTN\DamMinhLinh_ A03_BCDK1_ChuongII_hoanthien_v2.docx")
doc = Document(path)


def find_paragraph(text):
    return next(p for p in doc.paragraphs if p.text.strip() == text)


chapter2 = find_paragraph("CHƯƠNG 2. MÔ HÌNH ĐỀ XUẤT")
chapter3 = find_paragraph("CHƯƠNG 3. KẾT QUẢ THỰC NGHIỆM VÀ THẢO LUẬN")

# Sửa riêng citation trong Chương II để không phá số nguồn [1]–[5] của Chương I.
node = chapter2._p.getnext()
while node is not None and node is not chapter3._p:
    nxt = node.getnext()
    if node.tag == qn("w:p"):
        text = "".join(node.itertext())
        replacements = {
            "các lệnh show port stats/xstats được DPDK mô tả chính thức [2].":
                "các lệnh show port stats/xstats được DPDK mô tả chính thức [6].",
            "tối đa số packet yêu cầu [3].": "tối đa số packet yêu cầu [7].",
            "bộ đặc trưng chuẩn hoặc tối ưu [4].": "bộ đặc trưng chuẩn hoặc tối ưu [8].",
        }
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            # Các đoạn này chỉ có một run; ghi lại và giữ style/định dạng đoạn.
            for child in list(node):
                if child.tag == qn("w:r"):
                    node.remove(child)
            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Times New Roman")
            fonts.set(qn("w:hAnsi"), "Times New Roman")
            fonts.set(qn("w:eastAsia"), "Times New Roman")
            rpr.append(fonts)
            size = OxmlElement("w:sz")
            size.set(qn("w:val"), "26")
            rpr.append(size)
            run.append(rpr)
            t = OxmlElement("w:t")
            t.text = new_text
            run.append(t)
            node.append(run)
    node = nxt

# Xóa sáu mục tham khảo bị thêm thừa dưới dạng paragraph thường.
for p in list(doc.paragraphs):
    stripped = p.text.strip()
    if stripped.startswith((
        "[1] DPDK Project, “Getting Started Guide",
        "[2] DPDK Project, “Testpmd Application",
        "[3] DPDK Project, “rte_eth_rx_burst",
        "[4] I. Sharafaldin,",
        "[5] M. Ghadermazi và cộng sự,",
        "[6] Network Packet Transformation Approaches",
    )):
        p._element.getparent().remove(p._element)

# Hợp nhất vào content control Bibliography có sẵn: giữ [1]–[5], thêm [6]–[8].
references = [
    '[1] DPDK Project, "Documentation," [Online]. Available: https://core.dpdk.org/doc/. [Accessed 13 07 2026].',
    '[2] S. Eltanbouly, J. Zakraoui, A. Al-Ali, A. Belhi, S. Rahme and A. Bouras, "Network Packet Transformation Approaches for Intrusion Detection Systems: A Survey," IEEE Access, vol. 13, pp. 107293-107312, 2025.',
    '[3] J. Ghadermazi, A. Shah and N. D. Bastian, "Towards Real-Time Network Intrusion Detection With Image-Based Sequential Packets Representation," IEEE Transactions on Big Data, vol. 11, no. 1, pp. 157-173, 2025.',
    '[4] Y. Mirsky, T. Doitshman, Y. Elovici and A. Shabtai, "Kitsune: An Ensemble of Autoencoders for Online Network Intrusion Detection," Network and Distributed System Security Symposium (NDSS), San Diego, California, USA, 2018.',
    '[5] T. Zoppi and A. Ceccarelli, "Prepare for Trouble and Make it Double! Supervised–Unsupervised Stacking for Anomaly-Based Intrusion Detection," Journal of Network and Computer Applications, vol. 189, Art. no. 103106, 2021.',
    '[6] DPDK Project, "Testpmd Application User Guide: Runtime Functions and Running the Application," [Online]. Available: https://doc.dpdk.org/guides/testpmd_app_ug/. [Accessed 28 07 2026].',
    '[7] DPDK Project, "rte_eth_rx_burst — Ethernet Device API," [Online]. Available: https://doc.dpdk.org/api/rte__ethdev_8h.html. [Accessed 28 07 2026].',
    '[8] I. Sharafaldin, A. H. Lashkari and A. A. Ghorbani, "Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization," in Proc. ICISSP, 2018.',
]

sdts = doc._element.xpath(".//w:sdt[w:sdtContent]")
target_sdt = max(sdts, key=lambda e: len(e.findall(".//" + qn("w:p"))))
# Chọn content sâu nhất chứa các paragraph bibliography thực tế.
nested = [e for e in target_sdt.findall(".//" + qn("w:sdt")) if e.find(qn("w:sdtContent")) is not None]
if nested:
    target_sdt = max(nested, key=lambda e: len(e.findall(".//" + qn("w:p"))))
content = target_sdt.find(qn("w:sdtContent"))
for child in list(content):
    content.remove(child)
for ref in references:
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "Bibliography")
    ppr.append(style)
    p.append(ppr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = ref
    r.append(t)
    p.append(r)
    content.append(p)

# Chú thích bảng phải dùng style có sẵn và nằm dưới bảng theo yêu cầu.
caption = find_paragraph("Bảng 2.1. Danh sách 54 đặc trưng của nids.flow_features.v1")
table = next(t for t in doc.tables if len(t.rows) == 55 and len(t.columns) == 4)
caption._p.getparent().remove(caption._p)
table._tbl.addnext(caption._p)
caption.style = doc.styles["Caption"]
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_before = Pt(3)
caption.paragraph_format.space_after = Pt(6)
for run in caption.runs:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)

doc.save(path)
