from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(r"E:\DATTTN\TTTN\DamMinhLinh_ A03_BCDK1_ChuongII_hoanthien.docx")
doc = Document(path)
for paragraph in doc.paragraphs:
    ppr = paragraph._p.pPr
    if ppr is not None:
        for node in list(ppr):
            if node.tag == qn("w:shd"):
                ppr.remove(node)
settings = doc.settings._element
for node in list(settings):
    if node.tag == qn("w:updateFields"):
        settings.remove(node)

def find(text):
    return next(p for p in doc.paragraphs if p.text.strip() == text)

def remove_empty_between(start, end):
    node = start._p.getnext()
    while node is not None and node is not end._p:
        nxt = node.getnext()
        if node.tag == qn("w:p") and not "".join(node.itertext()).strip():
            node.getparent().remove(node)
        node = nxt

remove_empty_between(
    find("CHƯƠNG 2. MÔ HÌNH ĐỀ XUẤT"),
    find("CHƯƠNG 3. KẾT QUẢ THỰC NGHIỆM VÀ THẢO LUẬN"),
)
remove_empty_between(find("TÀI LIỆU THAM KHẢO"), find("PHỤ LỤC"))
doc.save(path)
