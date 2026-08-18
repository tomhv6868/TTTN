from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

path = Path(r"E:\DATTTN\TTTN\DamMinhLinh_ A03_BCDK1_ChuongII_hoanthien_v2.docx")
doc = Document(path)

fixes = {
    "discover chỉ thu thập": "discover chỉ thu thập và ghi bằng chứng; preflight kiểm tra topology, toolchain, devbind, hugepage và IOMMU; apply mới cấp hugepage và binding; run gọi dpdk-testpmd với EAL/application arguments đã khóa; rollback phục hồi host. Công cụ kiểm tra receipt từ chối bằng chứng thiếu bước hoặc sai cấu hình. testpmd cung cấp chế độ nhận và thống kê port; các lệnh show port stats/xstats được DPDK mô tả chính thức [6].",
    "Trong ứng dụng cuối nids_dpdk_live": "Trong ứng dụng cuối nids_dpdk_live, DPDK EAL được khởi tạo trước; port được cấu hình một RX queue, mempool mbuf được tạo, promiscuous mode được bật và vòng lặp gọi rte_eth_rx_burst. API DPDK mô tả hàm này là thao tác đọc các RX descriptor đã hoàn tất và trả về tối đa số packet yêu cầu [7]. Mỗi mbuf được chuyển sang DpdkAdapter; lỗi adapter, lỗi ingest, imissed và rx_nombuf được thống kê để phân biệt packet không hợp lệ, lỗi pipeline và drop ở port.",
    "Bộ 54 đặc trưng nids.flow_features.v1": "Bộ 54 đặc trưng nids.flow_features.v1 là hợp đồng nội bộ của dự án. Nó được điều chỉnh từ các nhóm đặc trưng CICFlowMeter nhưng chỉ dùng prefix 3, 5, 7 hoặc 9 packet, không dùng packet tương lai và không chờ flow kết thúc. Không có bài báo nào trong workspace công bố đúng bộ 54 này; vì vậy không được mô tả nó là bộ đặc trưng chuẩn hoặc tối ưu [8].",
}

for p in doc.paragraphs:
    for prefix, text in fixes.items():
        if p.text.startswith(prefix):
            p.text = text
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(13)

for p in doc.paragraphs:
    if p.text.strip() == "Bảng 2.1. Danh sách 54 đặc trưng của nids.flow_features.v1":
        for run in p.runs:
            run.bold = None
            run.italic = None

doc.save(path)
